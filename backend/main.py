import motor.motor_asyncio
import nltk
import docx
import os
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from dotenv import dotenv_values
from passlib.context import CryptContext

from fastapi import FastAPI, HTTPException, status, File, UploadFile
from fastapi import Body
from fastapi.responses import FileResponse
from fastapi.encoders import jsonable_encoder
from models import (UserCreate, UserLogin, ResponseModel, Token, TokenData)
from fastapi.middleware.cors import CORSMiddleware
from datetime import datetime, timedelta
from jose import JWTError, jwt
from llama_index.llms.ollama import Ollama
from docx import Document as DocxDocument

config = dotenv_values(".env")

client = motor.motor_asyncio.AsyncIOMotorClient(config["MONGODB_URI"])
db_name = config["DB_Name"]

database = client[db_name]

user_collection = database.get_collection("users_collection")

'''Creating a quick helper function for parsing the results from a db
   query into a python dict'''

def user_helper(user) -> dict:
    return {
        "name": user["name"],
        "email": user["email"],
        "password": user["password"]
    }

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def hash(password): 
    return pwd_context.hash(password, salt="a"*21 + "e")

def create_access_token(data: dict):
    to_encode = data.copy()  
    expire = datetime.now() + timedelta(minutes=1)
    to_encode.update({"exp": expire})  
    encoded_jwt = jwt.encode(to_encode, config["SECRET_KEY"], algorithm=config["ALGORITHM"])
    return encoded_jwt

# Add a new user into the db
async def add_user(user_data: dict) -> dict:  
    user = await user_collection.insert_one(user_data)
    new_user = await user_collection.find_one({"_id": user.inserted_id})
    return user_helper(new_user)

# Find a user from the db
async def find_user(user_data: dict) -> dict:
    user = await user_collection.find_one(user_data)
    found_user = await user_collection.find_one({"email": user["email"]})
    return user_helper(found_user)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/addUser", response_description="User data added into the database")
async def add_user_data(user: UserCreate = Body(...)):
    user = jsonable_encoder(user)
    hashed_pwd = hash(user["password"])
    user["password"] = hashed_pwd 
    new_user = await add_user(user)
    return ResponseModel(new_user, "User added successfully.")

@app.post('/findUser', response_description="Finding specific user data from the database")
async def find_user_data(user: UserLogin = Body(...)):
    user = jsonable_encoder(user)
    hashed_pwd = pwd_context.hash(user["password"], salt="a"*21 + "e")
    user["password"] = hashed_pwd
    found_user = await find_user(user)
    if not found_user:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Invalid credentails")
    
    # Creating and access token using the user's email as the payload
    user_data = {
        "email": found_user["email"]
    }
    access_token = create_access_token(data=user_data)
    print("access_token: ", access_token)
    return {"access_token": access_token, "token_type": "bearer"}

@app.post('/upload', response_description="Handling the file sent from the frontend")
async def upload(myFile: UploadFile = File(...)):
        contents = await myFile.read()

        #Saving to a temp file
        with open("temp.docx", "wb") as f:
            f.write(contents)

        #Parsing docx to extract text
        doc = DocxDocument("temp.docx")
        full_text = "\n".join([para.text for para in doc.paragraphs])

        print("File Name: ", myFile.filename)
        
        print("Full text: ", full_text) 

        # Tokenizing the text
        stopWords = set(stopwords.words("english"))
        words = word_tokenize(full_text)

        # Creating a frequency table to keep the score of each word
        freqTable = dict()
        for word in words:
            word = word.lower()
            if word in stopWords:
                 continue
            if word in freqTable:
                freqTable[word] += 1
            else:
                freqTable[word] = 1
            
        # Creating a dictionary to keep the score of each sentence
        sentences = sent_tokenize(full_text)
        sentenceValue = dict()

        for sentence in sentences:
            for word, freq in freqTable.items():
                if word in sentence.lower():
                    if sentence in sentenceValue:
                        sentenceValue[sentence] += freq
                    else:
                        sentenceValue[sentence] = freq
        
        sumValues = 0
        for sentence in sentenceValue:
            sumValues += sentenceValue[sentence]
        
        # Average value of a sentence from the original text
        average = int(sumValues / len(sentenceValue))

        # Writing the summary to a docx file
        doc = docx.Document()
        # Adding a paragraph to the document
        p = doc.add_paragraph()
        # Adding some formatting to the paragraph
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = 0

        # Storing sentences into our summary or docx file
        summary = ""
        for sentence in sentences:
            if (sentence in sentenceValue) and (sentenceValue[sentence] > (1.2 * average)):
                summary += " " + sentence
                
                # Adding a run to the paragraph
                p.add_run(sentence)

        print("summary: ", summary)

        # Saving the document
        doc.save("summary.docx")

@app.get("/download", response_description="Returning the summary docx file")
def download_file():
    # Getting the current working directory
    cwd = os.getcwd()
    file_path = cwd + "/summary.docx"
    return FileResponse(path=file_path, filename="summary.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        




